"""
DOCX文件解析器
使用python-docx库读取Word文档
"""

from docx import Document


def read_docx_file(file_path: str) -> str:
    """
    读取DOCX文件，返回纯文本内容
    段落之间用换行符分隔
    """
    doc = Document(file_path)
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)
    return "\n".join(paragraphs)


def read_docx_from_bytes(file_bytes: bytes) -> str:
    """从字节数据读取DOCX内容"""
    import io
    doc = Document(io.BytesIO(file_bytes))
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)
    return "\n".join(paragraphs)
